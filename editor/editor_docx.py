import os
import time
from docx import Document
from processamento.segmentador import segmentar_em_blocos
from processamento.revisor_llm import revisar_blocos_em_lote
from utils.config import AUTHOR
from utils.logger import LoggerProcesso
from transformers.tokenization_utils_fast import PreTrainedTokenizerFast
from modelo.carregador import tokenizer


def contar_tokens(blocos: list, tokenizer: PreTrainedTokenizerFast) -> int:
    """
    Soma o total de tokens dos blocos fornecidos.

    Args:
        blocos (list): Lista de blocos de texto.
        tokenizer (PreTrainedTokenizerFast): Tokenizador usado pelo modelo LLM.

    Returns:
        int: Total de tokens nas entradas.
    """
    return sum(len(tokenizer(bloco).input_ids) for bloco in blocos if isinstance(bloco, str) and bloco.strip())


def revisar_docx_otimizado(nome_arquivo: str):
    """
    Função principal que carrega um arquivo .docx com capítulos de uma webnovel,
    segmenta o texto em blocos, envia cada bloco para revisão por LLM e salva o resultado final revisado.

    Também gera log com:
    - tempo por capítulo,
    - blocos e tokens (entrada e saída),
    - quantidade de erros (fallbacks),
    - tempo total de execução.

    Args:
        nome_arquivo (str): Nome do arquivo .docx na pasta 'dados/entrada'.

    Returns:
        None. Salva documento revisado em 'dados/saida' e log em 'dados/logs'.
    """

    # Define caminhos
    caminho_entrada = os.path.join("dados", "entrada", nome_arquivo)
    nome_base = os.path.splitext(nome_arquivo)[0]
    caminho_saida = os.path.join("dados", "saida", f"{nome_base}_revisado.docx")

    # Carrega documento original
    doc = Document(caminho_entrada)

    # Separa capítulos com base nos títulos
    capitulos = []
    titulo_atual = None
    buffer = []

    for par in doc.paragraphs:
        if par.style and par.style.name in ("Heading 1", "Heading 2"):
            if titulo_atual and buffer:
                capitulos.append((titulo_atual, buffer))
                buffer = []
            titulo_atual = par.text.strip()
        else:
            texto = par.text.strip()
            if texto and texto != titulo_atual:
                buffer.append(texto)

    if titulo_atual and buffer:
        capitulos.append((titulo_atual, buffer))

    print(f"[📘] Total de capítulos identificados: {len(capitulos)}")

    # Cria novo doc e inicia logger
    novo_doc = Document()
    novo_doc.core_properties.author = AUTHOR
    logger = LoggerProcesso(nome_base)
    inicio_total = time.time()

    for i, (titulo, paragrafos) in enumerate(capitulos):
        print(f"[📖] {i+1}/{len(capitulos)}: {titulo}")
        inicio_capitulo = time.time()

        # Segmenta texto em blocos
        texto_capitulo = "\n".join(paragrafos).strip()
        blocos = segmentar_em_blocos(texto_capitulo, max_linhas=7)
        tokens_entrada = contar_tokens(blocos, tokenizer)

        # Revisão via LLM
        revisados, erros, tokens_saida, rev1, rev2, orig, recuperados_finais = revisar_blocos_em_lote(blocos, nome_base=nome_base, logger=logger)


        # Adiciona quebra de página e conteúdo revisado
        if i > 0:
            novo_doc.add_page_break()

        novo_doc.add_paragraph(titulo, style="Heading 1")

        for bloco in revisados:
            paragrafos = bloco.strip().split("\n")
            for par in paragrafos:
                novo_doc.add_paragraph(par.strip() if par.strip() else "")

        duracao = time.time() - inicio_capitulo
        logger.registrar_capitulo(
            titulo=titulo,
            blocos=len(blocos),
            tokens=tokens_entrada,
            erros=erros,
            duracao_segundos=duracao,
            tokens_saida=tokens_saida,
            rev1=rev1,
            rev2=rev2,
            orig=orig,
            recuperados=recuperados_finais,
        )
        print(f"[✅] Finalizado: {titulo} ({int(duracao // 60)}m {int(duracao % 60)}s)")


    # Salva arquivo final
    os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
    novo_doc.save(caminho_saida)
    logger.finalizar_log()

    # Tempo total
    duracao_total = time.time() - inicio_total
    horas = int(duracao_total // 3600)
    minutos = int((duracao_total % 3600) // 60)
    segundos = int(duracao_total % 60)

    print(f"\n[✓] Revisão concluída: {nome_arquivo}")
    print(f"[💾] Arquivo salvo em: {caminho_saida}")
    print(f"[⏱️] Tempo total: {horas}h {minutos}m {segundos}s")
    



